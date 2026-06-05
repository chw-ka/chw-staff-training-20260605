#!/usr/bin/env python3
"""Convert handouts/*.md to styled .docx; optional merged print edition with cover + TOC."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
HANDOUTS = ROOT / "handouts"
OUT_DIR = HANDOUTS / "docx"
MERGED_OUT = OUT_DIR / "CHW-Cursor-Training-Handouts-Merged.docx"

FONT = "Microsoft JhengHei"
FONT_CODE = "Consolas"
COLOR_HEADING = RGBColor(0x1E, 0x3A, 0x5F)
COLOR_ACCENT = RGBColor(0x25, 0x63, 0xEB)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
COLOR_TABLE_HEADER_BG = "E8F0FE"
COLOR_CODE_BG = "F1F5F9"
FOOTER_TEXT = "迦密聖道中學 · Cursor 進階培訓 · 2026-06-05"
TRAINING_DATE = "2026年6月5日"

SKIP_FILES: set[str] = set()

HANDOUT_SPECS: list[tuple[str, str]] = [
    ("README.md", "README.docx"),
    ("00-core-concepts-glossary.md", "00-core-concepts-glossary.docx"),
    ("08-appendix-安裝清單.md", "08-appendix-install-checklist.docx"),
    ("01-cursor-setup-guide.md", "01-cursor-setup-guide.docx"),
    ("02-prompt-cheatsheet.md", "02-prompt-cheatsheet.docx"),  # 標題：Prompt 實操指令速查
    ("07-static-site-publish.md", "07-static-site-publish.docx"),
    ("03-faq-hk-guide.md", "03-faq-hk-guide.docx"),
    ("05-api-key-application-guide.md", "05-api-key-application-guide.docx"),
    ("04-filesystem-mcp-guide.md", "04-filesystem-mcp-guide.docx"),
    ("06-google-drive-mcp-setup.md", "06-google-drive-mcp-setup.docx"),
    ("09-google-drive-self-study.md", "09-google-drive-self-study.docx"),
]

# Merged print order (excludes README index)
MERGE_SPECS: list[tuple[str, str]] = [
    ("00-core-concepts-glossary.md", "核心概念速查"),
    ("08-appendix-安裝清單.md", "附錄：課前安裝清單"),
    ("01-cursor-setup-guide.md", "Cursor 快速開始"),
    ("02-prompt-cheatsheet.md", "Prompt 實操指令速查 — 三個活動"),
    ("07-static-site-publish.md", "靜態網站發佈"),
    ("03-faq-hk-guide.md", "課後 FAQ"),
    ("05-api-key-application-guide.md", "API Key 申請指南"),
    ("04-filesystem-mcp-guide.md", "Filesystem MCP 試玩"),
    ("06-google-drive-mcp-setup.md", "Google Drive MCP 設定"),
    ("09-google-drive-self-study.md", "Google Drive 課後自學"),
]


def resolve_md_path(logical_name: str) -> Path | None:
    direct = HANDOUTS / logical_name
    if direct.exists():
        return direct
    prefix = logical_name.split("-", 1)[0] + "-"
    matches = sorted(HANDOUTS.glob(f"{prefix}*.md"))
    return matches[0] if matches else None


def extract_title(md_path: Path) -> str:
    text = md_path.read_text(encoding="utf-8")
    for line in text.splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return md_path.stem


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def set_run_font(run, name: str = FONT, size_pt: float = 11, bold: bool = False, color=None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element.rPr
    if r is not None:
        r.rFonts.set(qn("w:eastAsia"), name)
    else:
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_toc_field(doc: Document, heading: str = "目錄") -> None:
    doc.add_heading(heading, level=1)
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "（請在 Word 中對此目錄按右鍵 →「更新功能變數」）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(fld_text)
    r._r.append(fld_end)
    set_run_font(r, size_pt=10, color=COLOR_MUTED)
    doc.add_paragraph()


def add_print_cover(doc: Document) -> None:
    """Full cover page for merged handbook."""
    for _ in range(6):
        doc.add_paragraph()

    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p_school, "1E3A5F")
    r = p_school.add_run("迦密聖道中學")
    set_run_font(r, size_pt=28, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    p_en = doc.add_paragraph()
    p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_en.add_run("Carmel Holy Word Secondary School")
    set_run_font(r, size_pt=12, color=COLOR_MUTED)

    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("2026 Cursor 進階培訓")
    set_run_font(r, size_pt=22, bold=True, color=COLOR_HEADING)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("學員講義 · 合訂本")
    set_run_font(r, size_pt=16, color=COLOR_ACCENT)

    doc.add_paragraph()

    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p_tag, "E8F0FE")
    r = p_tag.add_run("從行政解放到高階自主開發 · Power Users 教職員")
    set_run_font(r, size_pt=11, color=COLOR_HEADING)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("培訓日期", TRAINING_DATE),
        ("建議時長", "90 分鐘"),
        ("版本", "2026-06-05"),
        ("內容", f"共 {len(MERGE_SPECS)} 章（概念、安裝、三活動、課後延伸）"),
    ]
    for i, (label, value) in enumerate(rows):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = value
        set_cell_shading(meta.rows[i].cells[0], COLOR_TABLE_HEADER_BG)
        for cell in meta.rows[i].cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    set_run_font(run, size_pt=10)

    doc.add_paragraph()
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_note.add_run("內部培訓資料 · 請勿外傳學生個人資料")
    set_run_font(r, size_pt=9, color=COLOR_MUTED)

    add_page_break(doc)


def add_cover_banner(doc: Document, doc_title: str) -> None:
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p1, "1E3A5F")
    r1 = p1.add_run("迦密聖道中學 · CHW")
    set_run_font(r1, size_pt=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p2, "E8F0FE")
    r2 = p2.add_run("2026 Cursor 進階培訓 · 學員講義")
    set_run_font(r2, size_pt=10, color=COLOR_HEADING)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(14)
    r3 = p3.add_run(doc_title)
    set_run_font(r3, size_pt=18, bold=True, color=COLOR_HEADING)
    doc.add_paragraph()


def configure_document(doc: Document, subtitle: str) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25
    pf.space_after = Pt(6)

    for level, size in [(1, 20), (2, 15), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_HEADING if level == 1 else COLOR_ACCENT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{FOOTER_TEXT}  |  {subtitle}")
    set_run_font(run, size_pt=9, color=COLOR_MUTED)


def add_inline_runs(paragraph, text: str, base_size: float = 11, monospace: bool = False) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, FONT_CODE if monospace else FONT, base_size)
        chunk = m.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_run_font(run, FONT, base_size, bold=True)
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            set_run_font(run, FONT_CODE, base_size - 0.5)
        elif chunk.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if link_m:
                label, url = link_m.groups()
                run = paragraph.add_run(f"{label} ({url})")
                set_run_font(run, FONT, base_size, color=COLOR_ACCENT)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, FONT_CODE if monospace else FONT, base_size)


def add_blockquote(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, "F8FAFC")
    text = " ".join(line.lstrip("> ").strip() for line in lines)
    add_inline_runs(p, text, base_size=10.5)
    for run in p.runs:
        run.italic = True


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(col_count):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, text, base_size=10)
            if i == 0:
                set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, COLOR_CODE_BG)
    text = "\n".join(lines)
    run = p.add_run(text)
    set_run_font(run, FONT_CODE, 9.5)
    doc.add_paragraph()


def add_list(doc: Document, items: list[tuple[str, str]]) -> None:
    for kind, text in items:
        style = "List Number" if kind == "ol" else "List Bullet"
        p = doc.add_paragraph(style=style)
        add_inline_runs(p, text)


def build_document_from_markdown(
    md_path: Path,
    *,
    include_banner: bool = True,
    section_label: str | None = None,
    skip_first_h1: bool = True,
) -> Document:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = extract_title(md_path)

    doc = Document()
    configure_document(doc, "學員講義合訂本" if section_label else md_path.name)

    if section_label:
        doc.add_heading(section_label, level=1)
        doc.add_paragraph()
    elif include_banner:
        add_cover_banner(doc, title)

    i = 0
    list_buffer: list[tuple[str, str]] = []
    skipped_h1 = False

    def flush_list() -> None:
        nonlocal list_buffer
        if list_buffer:
            add_list(doc, list_buffer)
            list_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_list()
            i += 1
            continue

        if stripped == "---":
            flush_list()
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("```"):
            flush_list()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_list()
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_separator(lines[i]):
                    table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        if stripped.startswith(">"):
            flush_list()
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i])
                i += 1
            add_blockquote(doc, quote_lines)
            continue

        hm = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if hm:
            flush_list()
            level = len(hm.group(1))
            heading_text = hm.group(2).strip()
            if level == 1 and skip_first_h1 and not skipped_h1:
                skipped_h1 = True
                i += 1
                continue
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        ol_m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        ul_m = re.match(r"^[-*]\s+(.+)$", stripped)
        if ol_m:
            list_buffer.append(("ol", ol_m.group(2)))
            i += 1
            continue
        if ul_m:
            list_buffer.append(("ul", ul_m.group(1)))
            i += 1
            continue

        flush_list()
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    flush_list()
    return doc


def convert_markdown(md_path: Path, out_path: Path) -> None:
    doc = build_document_from_markdown(md_path, include_banner=True)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  OK  {out_path.relative_to(ROOT)}")


def build_merged_docx(out_path: Path = MERGED_OUT) -> None:
    from docxcompose.composer import Composer

    print(f"Building merged handbook -> {out_path.relative_to(ROOT)}")

    master = Document()
    configure_document(master, "學員講義合訂本")
    add_print_cover(master)
    add_toc_field(master)
    add_page_break(master)

    composer = Composer(master)
    part = 0
    for logical_md, section_title in MERGE_SPECS:
        md = resolve_md_path(logical_md)
        if md is None:
            print(f"  SKIP merge: {logical_md}")
            continue
        part += 1
        label = f"第 {part} 章　{section_title}"
        body = build_document_from_markdown(
            md,
            include_banner=False,
            section_label=label,
            skip_first_h1=True,
        )
        composer.append(body)
        print(f"  + {label}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(out_path))
    print(f"  OK  {out_path.relative_to(ROOT)}")
    print("  Tip: Open in Word -> right-click TOC -> Update Field (更新功能變數)")


def main() -> None:
    args = sys.argv[1:]
    merged_only = "--merged-only" in args
    separate_only = "--separate-only" in args

    if not merged_only:
        print(f"Converting {len(HANDOUT_SPECS)} handouts -> {OUT_DIR.relative_to(ROOT)}/")
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        used_stems: set[str] = set()
        for logical_md, out_name in HANDOUT_SPECS:
            md = resolve_md_path(logical_md)
            if md is None:
                print(f"  SKIP (missing) {logical_md}")
                continue
            used_stems.add(md.name)
            convert_markdown(md, OUT_DIR / out_name)
        for md in sorted(HANDOUTS.glob("*.md")):
            if md.name in used_stems or md.name in SKIP_FILES:
                continue
            convert_markdown(md, OUT_DIR / f"{md.stem}.docx")

    if not separate_only:
        build_merged_docx()

    print("Done.")


if __name__ == "__main__":
    main()
