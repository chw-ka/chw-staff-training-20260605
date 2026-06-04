#!/usr/bin/env python3
"""Generate sample agenda and last-year minutes as .docx (Chinese filenames)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def build_agenda() -> Path:
    out = ROOT / "議程_視藝科組會_20260528.docx"
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "視覺藝術科組會議 — 議程")

    meta = [
        ("會議日期", "2026年5月28日（星期三）"),
        ("時間", "16:00–17:05"),
        ("地點", "視藝室"),
        ("主席", "陳主任（視藝科主任）"),
        ("記錄", "林老師"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(f"{label}：").bold = True
        p.add_run(value)

    doc.add_paragraph()
    doc.add_heading("議程", level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["編號", "議項", "預計時間", "負責報告"]
    rows = [
        ["1", "下學期視藝科展覽安排（校庆 cross-media showcase）", "25 分鐘", "陳主任"],
        ["2", "中四選修課程及 split class 編排", "20 分鐘", "陳主任"],
        ["3", "AI 工具於評核功課的使用指引（科組補充規定）", "15 分鐘", "陳主任"],
        ["4", "其他事項", "5 分鐘", "—"],
    ]
    for c, h in enumerate(headers):
        table.rows[0].cells[c].text = h
        for p in table.rows[0].cells[c].paragraphs:
            for r in p.runs:
                r.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()
    doc.add_heading("會前閱讀", level=2)
    for item in [
        "教务处 AI 使用指引 draft（2026年5月版）",
        "上學年展覽材料預算表（參考）",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("預期產出", level=2)
    for item in [
        "確認展覽主題、三區安排及 submission deadline",
        "確定與 Timetable 組跟進 split class 的時間表",
        "科組 AI declaration form 起草人及下次會議批准流程",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.save(out)
    return out


def add_section(doc: Document, title: str, summary: str, resolutions: list[str], followups: list[tuple]) -> None:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.add_run("討論摘要：").bold = True
    doc.add_paragraph(summary)
    p = doc.add_paragraph()
    p.add_run("決議：").bold = True
    for i, res in enumerate(resolutions, 1):
        doc.add_paragraph(f"{i}. {res}", style="List Number")
    p = doc.add_paragraph()
    p.add_run("跟進：").bold = True
    if followups:
        t = doc.add_table(rows=1 + len(followups), cols=4)
        t.style = "Table Grid"
        for c, h in enumerate(["負責人", "事項", "截止日期", "狀態"]):
            t.rows[0].cells[c].text = h
        for r, row in enumerate(followups, start=1):
            for c, val in enumerate(row):
                t.rows[r].cells[c].text = val
    doc.add_paragraph()


def build_last_year_minutes() -> Path:
    out = ROOT / "會議紀錄_視藝科組_20250522_上學年.docx"
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "會議紀錄")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("（格式參考 — 供第三步跟版面、欄位用；內容來自今年逐字稿，唔係抄本檔）").italic = True

    fields = [
        ("會議名稱", "視覺藝術科組會議"),
        ("日期", "2025年5月22日（星期四）"),
        ("時間", "16:00–16:50"),
        ("地點", "視藝室"),
        ("主席", "陳主任"),
        ("記錄", "林老師"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}：").bold = True
        p.add_run(value)

    doc.add_paragraph()
    doc.add_heading("二、議程及討論摘要（節錄 — 與今年相關）", level=2)

    add_section(
        doc,
        "議程 1：學生作品展預備",
        "科組初步討論校庆前舉辦學生作品展，主題方向為「城市與記憶」構思，尚未定案三區細則。",
        ["本學年先完成主題構思，下學年開學後再定 submission 規格"],
        [
            ("陳主任", "向校長提交展覽構思書", "2025年6月", "已完成"),
            ("各老師", "收集學生作品意向", "2025年10月", "已完成"),
        ],
    )

    add_section(
        doc,
        "議程 2：中四 split class 初探",
        "2025年已有 split class 現象，但未達今年規模；科組要求 Timetable 組預留彈性。",
        ["原則：split class 盡量不超過 8 人（今年討論改為 10 人）"],
        [("林老師", "整理 split 名單模板", "2025年9月", "已完成 — 今年需更新名單")],
    )

    add_section(
        doc,
        "議程 3：AI 工具使用（首次討論）",
        "科組關注 AI 生成作品對評核的影響，待教务处發出全校指引後再定科組規定。",
        ["暫緩，待教务处指引"],
        [("陳主任", "留意教务处 draft", "2026年5月", "教务处已出 draft — 今年主動討論")],
    )

    doc.add_heading("三、格式參考說明", level=2)
    for item in [
        "本檔只供第三步跟足版面、欄位、語氣",
        "今年紀錄內容必須來自議程同逐字稿，不可抄本檔舊內容",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph()
    p = doc.add_paragraph("（上學年紀錄格式範例完）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out)
    return out


def main() -> None:
    agenda = build_agenda()
    minutes = build_last_year_minutes()
    print(f"Created: {agenda.name}")
    print(f"Created: {minutes.name}")


if __name__ == "__main__":
    main()
